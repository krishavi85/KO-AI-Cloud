from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".csv",
    ".json",
    ".py",
    ".js",
    ".jsx",
    ".ts",
    ".tsx",
    ".html",
    ".css",
    ".xml",
    ".yaml",
    ".yml",
}


def parse_txt(path: str) -> str:
    return Path(path).read_text(encoding="utf-8", errors="ignore")


def parse_pdf(path: str) -> str:
    reader = PdfReader(path)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def parse_docx(path: str) -> str:
    doc = DocxDocument(path)
    paragraphs = [paragraph.text for paragraph in doc.paragraphs]
    table_cells = []
    for table in doc.tables:
        for row in table.rows:
            table_cells.extend(cell.text for cell in row.cells)
    return "\n".join(paragraphs + table_cells)


def parse_file(path: str, content_type: str | None = None) -> str:
    suffix = Path(path).suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(path)
    if suffix == ".docx":
        return parse_docx(path)
    if suffix in TEXT_EXTENSIONS:
        return parse_txt(path)
    raise ValueError(f"Unsupported file type: {suffix or content_type or 'unknown'}")


def chunk_text(text: str, chunk_size: int = 1200, overlap: int = 200) -> list[str]:
    if chunk_size <= overlap:
        raise ValueError("chunk_size must be greater than overlap")

    clean = " ".join(text.split())
    if not clean:
        return []

    chunks = []
    start = 0
    while start < len(clean):
        end = min(start + chunk_size, len(clean))
        chunk = clean[start:end]
        if chunk.strip():
            chunks.append(chunk)
        if end == len(clean):
            break
        start = end - overlap
    return chunks
